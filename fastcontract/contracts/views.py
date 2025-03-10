from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse, JsonResponse, FileResponse
from .models import ContractTemplate, ContractInstance
from .forms import ContractTemplateForm, ContractInstanceForm, SendContractForm
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
import json


def upload_template(request):
    if request.method == 'POST':
        form = ContractTemplateForm(request.POST, request.FILES)
        if form.is_valid():
            template = form.save()
            variables = template.extract_variables()
            return HttpResponse(f"Template uploaded successfully. Extracted variables: {variables}")
    else:
        form = ContractTemplateForm()
    return render(request, 'contracts/upload_template.html', {'form': form, 'css_file': 'https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css'})

def create_contract_instance(request, template_id):
    contract_template = get_object_or_404(ContractTemplate, id=template_id)
    if request.method == 'POST':
        data = json.loads(request.body)
        instances = data['instances']
        for instance in instances:
            variables = instance['variables']
            values = instance['values']
            contract_instance = ContractInstance(template=contract_template, variables=dict(zip(variables, values)))
            contract_instance.save()
        return JsonResponse({'success': True})
    return JsonResponse({'success': False})

def list_templates(request):
    if request.method == 'POST':
        form = ContractTemplateForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
    else:
        form = ContractTemplateForm()
    templates = ContractTemplate.objects.all()
    return render(request, 'contracts/list_templates.html', {'templates': templates, 'form': form, 'css_file': 'https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css'})

# def send_contract(request, instance_id):
#     instance = ContractInstance.objects.get(id=instance_id)
#     if request.method == 'POST':
#         form = SendContractForm(request.POST)
#         if form.is_valid():
#             # Logic to send contract for signature
#             return HttpResponse(f"Contract sent for signature successfully.")
#     else:
#         form = SendContractForm()
#     return render(request, 'contracts/send_contract.html', {'form': form, 'instance': instance, 'css_file': 'https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css'})

def display_template_pdf(request, template_id):
    template = get_object_or_404(ContractTemplate, id=template_id)
    document = Document(template.template_file.path)
    
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    y = height - 40
    for para in document.paragraphs:
        p.drawString(40, y, para.text)
        y -= 20
        if y < 40:
            p.showPage()
            y = height - 40

    p.save()
    buffer.seek(0)
    pdf = buffer.getvalue()
    buffer.close()

    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="{template.name}.pdf"'
    return response

def download_template(request, template_id):
    template = get_object_or_404(ContractTemplate, id=template_id)
    response = FileResponse(template.template_file.open(), as_attachment=True, filename=template.template_file.name)
    return response

def contract_instances(request):
    contract_instances = ContractInstance.objects.all()
    context = {
        'contract_instances': [
            {
                'template_name': instance.template.name,
                'variables': list(instance.variables.keys()),
                'values': list(instance.variables.values()),
                'id':instance.id
            }
            for instance in contract_instances
        ]
    }
    return render(request, 'contracts/contract_instances.html', context)

def download_contract_instance(request, instance_id):
    instance = get_object_or_404(ContractInstance, id=instance_id)
    document = Document(instance.template.template_file.path)
    
    # Replace variables in the document
    for variable, value in instance.variables.items():
        placeholder = f"<<% {variable} %>>"
        for paragraph in document.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Create a response with the document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={instance.template.name}.docx'
    document.save(response)
    return response

def excel_like_table(request):
    return render(request, 'contracts/excel_like_table.html')